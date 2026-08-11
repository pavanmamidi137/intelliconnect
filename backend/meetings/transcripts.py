"""Transcript parsing for supported meeting transcript formats.

Supported: TXT, PDF, DOCX, SRT, VTT. Unsupported formats raise a
FileValidationError with a friendly message.
"""

import re
from pathlib import Path

from config.exceptions import FileValidationError

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".srt", ".vtt"}

# A speaker turn: "Name: text" or "Name said: text" (labels like
# "Speaker 1", "Ravi Kumar", "Interviewer" etc.).
_SPEAKER_LINE_RE = re.compile(
    r"^\s*(?P<speaker>[A-Za-z][A-Za-z0-9 .'’-]{0,48}?)\s*: \s*(?P<text>.+?)\s*$"
)
_KNOWN_NON_SPEAKER_WORDS = {
    "note", "notes", "agenda", "summary", "transcript", "meeting", "action",
    "actions", "attendees", "participants", "discussion", "decision",
    "decisions", "follow-up", "followup", "status", "time", "date",
    "location", "minutes", "next", "item", "items", "http", "www",
}


def normalize_name(value: str) -> str:
    """Lowercase, strip titles and punctuation for fuzzy name matching."""
    name = re.sub(r"^\s*(mr|mrs|ms|miss|dr|prof|sir|ma'am|madam)\s+", "",
                  value.strip().lower())
    return re.sub(r"[^a-z0-9 ]", "", name).strip()


def parse_transcript_turns(content: str):
    """Split transcript text into speaker turns.

    Returns (turns, has_speakers) where each turn is
    {"speaker": str, "text": str}. Lines without a "Name:" label are
    appended to the previous turn's text; when no labelled speakers exist
    at all, the whole text is returned as a single anonymous turn.
    """
    turns = []
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = _SPEAKER_LINE_RE.match(line)
        if match and normalize_name(match.group("speaker")) not in _KNOWN_NON_SPEAKER_WORDS:
            speaker = match.group("speaker").strip()
            text = match.group("text").strip()
            turns.append({"speaker": speaker, "text": text})
        else:
            if turns:
                turns[-1]["text"] = f"{turns[-1]['text']} {line}".strip()
            else:
                turns.append({"speaker": "", "text": line})

    # Collapse consecutive turns from the same speaker.
    collapsed = []
    for turn in turns:
        if collapsed and collapsed[-1]["speaker"].lower() == turn["speaker"].lower():
            collapsed[-1]["text"] = f"{collapsed[-1]['text']}\n{turn['text']}"
        else:
            collapsed.append(dict(turn))

    has_speakers = any(turn["speaker"] for turn in collapsed)
    return collapsed, has_speakers


def is_supported_transcript(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_transcript_text(filename: str, content: bytes) -> str:
    """Return plain text from a transcript file of any supported format."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise FileValidationError(
            "This file type isn't supported. Upload a TXT, PDF, DOCX, SRT, or VTT transcript."
        )

    try:
        if ext == ".txt":
            text = _decode(content)
        elif ext == ".pdf":
            text = _extract_pdf(content)
        elif ext == ".docx":
            text = _extract_docx(content)
        elif ext in (".srt", ".vtt"):
            text = _strip_caption_timestamps(_decode(content))
        else:  # pragma: no cover - guarded above
            text = ""
    except FileValidationError:
        raise
    except Exception as exc:
        raise FileValidationError(
            f"We couldn't read this file. Please make sure it's a valid {ext[1:].upper()} transcript."
        ) from exc

    text = text.strip()
    if not text:
        raise FileValidationError(
            "This file appears to be empty. Please upload a transcript with content."
        )
    return text


def _decode(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise FileValidationError("We couldn't read this file's encoding.")


def _extract_pdf(content: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    from io import BytesIO

    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


_TIMESTAMP_RE = re.compile(
    r"(?m)^(\d{1,2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{1,2}:\d{2}:\d{2}[,.]\d{3}"
    r"|WEBVTT.*$|\d{1,2}:\d{2}\.\d{3}\s*-->\s*\d{1,2}:\d{2}\.\d{3}.*$)"
)


def _strip_caption_timestamps(text: str) -> str:
    """Remove SRT/VTT cue indexes, timestamps and formatting tags."""
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.isdigit():
            continue
        if "-->" in stripped or stripped.upper().startswith("WEBVTT"):
            continue
        stripped = re.sub(r"<[^>]+>", "", stripped)
        if stripped:
            lines.append(stripped)
    return "\n".join(lines)
